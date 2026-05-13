from __future__ import annotations

import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from src.ingestion.parsers.base import DocumentParser, ParserError
from src.ingestion.parsers.asset_utils import guess_mime_type, resolve_ooxml_target, write_asset_file
from src.models.schemas import ParsedDocument, Section
from src.models.schemas import Asset

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PACKAGE_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


class DocxParser(DocumentParser):
    """DOCX parser using paragraph extraction plus embedded media extraction."""

    file_type = "docx"

    def parse(self, path: Path, asset_output_dir: Path | None = None) -> ParsedDocument:
        try:
            from docx import Document  # type: ignore[import-not-found]
        except Exception as exc:
            raise ParserError("DOCX parsing requires python-docx.") from exc

        try:
            document = Document(str(path))
        except Exception as exc:
            raise ParserError(f"Failed to open DOCX {path}: {exc}") from exc

        doc_id = self.doc_id(path)
        title = self.title_from_path(path)
        assets = _extract_embedded_images(path, doc_id, asset_output_dir)
        sections: list[Section] = []
        current_title = title
        current_level = 1
        current_lines: list[str] = []
        section_index = 0
        all_lines: list[str] = []

        def flush_section() -> None:
            nonlocal section_index, current_lines
            text = "\n".join(line for line in current_lines if line.strip()).strip()
            if not text and not current_title:
                return
            section_index += 1
            sections.append(
                Section(
                    section_id=f"{doc_id}-section-{section_index}",
                    title=current_title or f"Section {section_index}",
                    level=current_level,
                    text=text,
                    metadata={"parser": "docx"},
                )
            )
            current_lines = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            all_lines.append(text)
            style_name = paragraph.style.name if paragraph.style is not None else ""
            heading_level = _heading_level(style_name)
            if heading_level is not None:
                flush_section()
                current_title = text
                current_level = heading_level
            else:
                current_lines.append(text)

        flush_section()
        body = "\n".join(all_lines).strip()
        return ParsedDocument(
            doc_id=doc_id,
            file_path=str(path.resolve()),
            file_type=self.file_type,
            title=title,
            text=body,
            sections=sections,
            assets=assets,
            metadata={"paragraph_count": len(document.paragraphs)},
        )


def _heading_level(style_name: str) -> int | None:
    normalized = style_name.strip().lower()
    if not normalized.startswith("heading"):
        return None
    parts = normalized.split()
    if len(parts) > 1 and parts[-1].isdigit():
        return max(1, int(parts[-1]))
    return 1


def _extract_embedded_images(path: Path, doc_id: str, asset_output_dir: Path | None) -> list[Asset]:
    try:
        archive = zipfile.ZipFile(path)
    except Exception as exc:
        raise ParserError(f"Failed to inspect embedded DOCX assets in {path}: {exc}") from exc

    with archive:
        media_parts = {
            name: archive.read(name)
            for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        }
        document_part = "word/document.xml"
        rels_part = "word/_rels/document.xml.rels"
        if document_part not in archive.namelist():
            return []
        rels = _read_relationships(archive, rels_part, document_part)
        root = ET.fromstring(archive.read(document_part))
        assets: list[Asset] = []
        for image_index, drawing in enumerate(root.findall(f".//{{{W_NS}}}drawing"), start=1):
            blip = drawing.find(f".//{{{A_NS}}}blip")
            if blip is None:
                continue
            embed = blip.attrib.get(f"{{{R_NS}}}embed", "")
            target = rels.get(embed, "")
            data = media_parts.get(target)
            if not data:
                continue
            doc_pr = drawing.find(f".//{{{WP_NS}}}docPr")
            extent = drawing.find(f".//{{{WP_NS}}}extent")
            asset_id = f"{doc_id}-image-{image_index:04d}"
            source_name = Path(target).name or f"image-{image_index}.bin"
            stored_path, sha256 = write_asset_file(asset_output_dir, doc_id, asset_id, source_name, data)
            assets.append(
                Asset(
                    asset_id=asset_id,
                    kind="image",
                    source_file=str(path.resolve()),
                    page_or_slide=None,
                    path=stored_path,
                    caption=(doc_pr.attrib.get("descr", "") if doc_pr is not None else "") or (
                        doc_pr.attrib.get("name", "") if doc_pr is not None else ""
                    ),
                    mime_type=guess_mime_type(source_name),
                    sha256=sha256,
                    width=_optional_int(extent.attrib.get("cx") if extent is not None else None),
                    height=_optional_int(extent.attrib.get("cy") if extent is not None else None),
                    metadata={
                        "relationship_id": embed,
                        "package_path": target,
                        "shape_name": doc_pr.attrib.get("name", "") if doc_pr is not None else "",
                    },
                )
            )
        return assets


def _read_relationships(archive: zipfile.ZipFile, rels_part: str, base_part: str) -> dict[str, str]:
    if rels_part not in archive.namelist():
        return {}
    root = ET.fromstring(archive.read(rels_part))
    rels: dict[str, str] = {}
    for rel in root.findall(f".//{{{PACKAGE_NS}}}Relationship"):
        rel_id = rel.attrib.get("Id", "")
        target = rel.attrib.get("Target", "")
        if rel_id and target:
            rels[rel_id] = resolve_ooxml_target(base_part, target)
    return rels


def _optional_int(value: str | None) -> int | None:
    if value is None:
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None
